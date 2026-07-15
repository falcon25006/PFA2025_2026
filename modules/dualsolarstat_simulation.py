# -*- coding: utf-8 -*-
"""
DualSolarStat3 — Application complète en un seul fichier
============================================================
Ce fichier unique contient :
  A) LE MOTEUR DE CALCUL (chaînes de Markov, simulation Monte-Carlo,
     chargement de données CSV, gestion d'erreurs)
  B) L'INTERFACE GRAPHIQUE (CustomTkinter — apparence moderne)
  C) L'EXPORT DES RÉSULTATS (.txt et .csv)

Lancement :
    python3 Dualsolarstat3.py

Dépendances :
    pip install customtkinter numpy --break-system-packages

Structure du fichier (recherchez ces bannières pour naviguer) :
    ==== SECTION A : MOTEUR DE CALCUL ====
    ==== SECTION B : INTERFACE GRAPHIQUE ====
    ==== SECTION C : POINT D'ENTREE ====
"""

import csv
import os
import queue
import tempfile
import threading
import uuid
import numpy as np

import customtkinter as ctk
from docx import Document
from docx.shared import Inches
from tkinter import filedialog, messagebox
from matplotlib.figure import Figure
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from mpl_toolkits.mplot3d import Axes3D  # noqa: F401


# ============================================================================
# ==== SECTION A : MOTEUR DE CALCUL =========================================
# ============================================================================
# (Logique identique à dualsolarstat_core.py, testée indépendamment de
# l'interface graphique — aucune ligne de cette section ne dépend de Tk.)


class DualSolarStatError(Exception):
    """
    Exception dédiée au projet. Toute erreur prévisible (fichier absent,
    colonne manquante, valeur invalide, liste trop courte...) lève CETTE
    exception avec un message en français directement affichable dans
    l'interface — jamais une trace Python brute à l'utilisateur.
    """
    pass


# --- Valeurs par défaut de la "structure test" (cas Mairie de Boulsa) -----

DEFAUTS = {
    "e_jour_base_wh": 108900.0,
    "eta_systeme": 0.90,
    "pc_panneau_wc": 600.0,
    "capacite_batterie_ah": 300.0,
    "tension_bus_v": 51.2,
    "profondeur_decharge": 0.90,
    "n_panneaux_test": 28,
    "n_batteries_test": 4,
    "cs_etats": [0.60, 0.75, 0.90],
    "cs_matrice": [
        [0.55, 0.35, 0.10],
        [0.25, 0.50, 0.25],
        [0.10, 0.35, 0.55],
    ],
    "psh_etats": [6.5, 5.0, 2.5],
    "psh_matrice": [
        [0.75, 0.20, 0.05],
        [0.30, 0.55, 0.15],
        [0.15, 0.35, 0.50],
    ],
}

TYPES_MENAGE = {
    "faible":  {"lolp_cible": 0.10, "libelle": "Consommation faible (LOLP 10%)"},
    "moyenne": {"lolp_cible": 0.05, "libelle": "Consommation moyenne (LOLP 5%)"},
    "elevee":  {"lolp_cible": 0.02, "libelle": "Consommation élevée (LOLP 2%)"},
}
TYPE_MENAGE_PAR_DEFAUT = "moyenne"


# --- Chargement de données CSV (avec gestion d'erreurs complète) ----------

def _lire_csv_colonne_numerique(chemin_fichier, nom_colonne, description):
    """Lit une colonne numérique dans un CSV. Lève DualSolarStatError si problème."""
    if not os.path.isfile(chemin_fichier):
        raise DualSolarStatError(
            f"Fichier introuvable pour les données de {description} : « {chemin_fichier} »."
        )
    valeurs = []
    try:
        with open(chemin_fichier, newline="", encoding="utf-8-sig") as f:
            lecteur = csv.DictReader(f)
            if lecteur.fieldnames is None or nom_colonne not in lecteur.fieldnames:
                raise DualSolarStatError(
                    f"La colonne « {nom_colonne} » est absente du fichier {description}. "
                    f"Colonnes trouvées : {lecteur.fieldnames}."
                )
            for i, ligne in enumerate(lecteur, start=2):
                brut = (ligne.get(nom_colonne) or "").strip().replace(",", ".")
                if brut == "":
                    continue
                try:
                    valeurs.append(float(brut))
                except ValueError:
                    raise DualSolarStatError(
                        f"Valeur non numérique ligne {i} du fichier {description} "
                        f"(colonne « {nom_colonne} ») : « {brut} »."
                    )
    except UnicodeDecodeError:
        raise DualSolarStatError(f"Le fichier {description} n'est pas un CSV texte lisible (encodage invalide).")
    except csv.Error as e:
        raise DualSolarStatError(f"Erreur de lecture du fichier {description} : {e}")

    if len(valeurs) < 2:
        raise DualSolarStatError(
            f"Le fichier {description} contient trop peu de valeurs exploitables "
            f"({len(valeurs)}) : il en faut au moins 2."
        )
    if any(v < 0 for v in valeurs):
        raise DualSolarStatError(f"Le fichier {description} contient des valeurs négatives.")
    return valeurs


def charger_consommation_mensuelle(chemin_fichier, nom_colonne="consommation_wh"):
    """
    Charge une liste de consommations MENSUELLES TOTALES (Wh). Longueur
    variable (2 mois minimum, idéalement 12+). ATTENTION : total mensuel
    réel, pas une moyenne journalière (voir construire_parametres()).
    """
    return _lire_csv_colonne_numerique(chemin_fichier, nom_colonne, "consommation mensuelle (SONABEL)")


def charger_ensoleillement_journalier(chemin_fichier, nom_colonne="psh_heures"):
    """Charge une liste d'ensoleillement journalier (PSH, en heures). Longueur variable."""
    valeurs = _lire_csv_colonne_numerique(chemin_fichier, nom_colonne, "ensoleillement journalier (météo)")
    if any(v > 13 for v in valeurs):
        raise DualSolarStatError(
            "Le fichier d'ensoleillement contient des valeurs de PSH supérieures à 13 heures : "
            "vérifiez l'unité utilisée (heures attendues)."
        )
    return valeurs


def charger_catalogue_materiel(chemin_fichier):
    """
    Charge le catalogue matériel (panneaux/batteries/onduleurs + prix) depuis
    un CSV fourni par l'utilisateur. Aucune valeur par défaut : dépend du
    pays/fournisseur.
    Format : type,nom,caracteristique,prix
        - panneau  : caracteristique = puissance crête (Wc)
        - batterie : caracteristique = capacité (Ah)
        - onduleur : caracteristique = puissance nominale (W)
    """
    if not os.path.isfile(chemin_fichier):
        raise DualSolarStatError(
            f"Fichier catalogue introuvable : « {chemin_fichier} ». Le catalogue matériel doit "
            f"être fourni par l'utilisateur (aucune valeur par défaut n'existe pour ce fichier)."
        )
    panneaux, batteries, onduleurs = [], [], []
    colonnes_requises = {"type", "nom", "caracteristique", "prix"}
    with open(chemin_fichier, newline="", encoding="utf-8-sig") as f:
        lecteur = csv.DictReader(f)
        if lecteur.fieldnames is None or not colonnes_requises.issubset(set(lecteur.fieldnames)):
            raise DualSolarStatError(
                f"Le catalogue doit contenir les colonnes {sorted(colonnes_requises)}. "
                f"Colonnes trouvées : {lecteur.fieldnames}."
            )
        for i, ligne in enumerate(lecteur, start=2):
            type_item = (ligne.get("type") or "").strip().lower()
            if type_item not in ("panneau", "batterie", "onduleur"):
                raise DualSolarStatError(
                    f"Ligne {i} du catalogue : le type doit être « panneau », « batterie » ou « onduleur »."
                )
            try:
                caracteristique = float((ligne.get("caracteristique") or "").replace(",", "."))
                prix = float((ligne.get("prix") or "").replace(",", "."))
            except ValueError:
                raise DualSolarStatError(f"Ligne {i} du catalogue : caractéristique ou prix non numérique.")
            if caracteristique <= 0 or prix <= 0:
                raise DualSolarStatError(f"Ligne {i} du catalogue : caractéristique et prix doivent être positifs.")
            item = {"nom": ligne.get("nom", "").strip(), "caracteristique": caracteristique, "prix": prix}
            if type_item == "panneau":
                panneaux.append(item)
            elif type_item == "batterie":
                batteries.append(item)
            else:
                onduleurs.append(item)

    if not panneaux or not batteries or not onduleurs:
        raise DualSolarStatError(
            "Le catalogue doit contenir au moins un panneau, une batterie ET un onduleur."
        )
    return {"panneaux": panneaux, "batteries": batteries, "onduleurs": onduleurs}


# Marge de sécurité appliquée à la puissance crête installée pour dimensionner
# l'onduleur (même convention que la méthode traditionnelle du technicien :
# cf. tableau "Onduleur ... Marge de sécurité fixe ~25%" du cas Boulsa).
# Limite assumée : faute de courbe de charge horaire, la puissance de
# pointe des CHARGES n'est pas modélisée séparément ici ; on dimensionne
# l'onduleur pour qu'il puisse convertir toute la puissance crête PV
# installée, ce qui est l'approche standard en l'absence de profil de
# puissance instantanée du site.
MARGE_SECURITE_ONDULEUR = 0.25
SURFACE_M2_PAR_KWC = 3.5


def estimer_surface_installation(n_panneaux, pc_panneau_wc):
    """Estime la surface nécessaire en m² pour une installation donnée."""
    return (n_panneaux * pc_panneau_wc / 1000.0) * SURFACE_M2_PAR_KWC


def dimensionner_onduleur(puissance_crete_installee_wc, onduleur_catalogue):
    """
    Détermine le nombre d'onduleurs nécessaires (et leur coût) pour couvrir
    la puissance crête installée, avec la marge de sécurité standard.

    Retour
    ------
    dict {"n_onduleurs": int, "puissance_necessaire_w": float, "cout": float}
    """
    puissance_necessaire = puissance_crete_installee_wc * (1 + MARGE_SECURITE_ONDULEUR)
    n_onduleurs = max(1, int(np.ceil(puissance_necessaire / onduleur_catalogue["caracteristique"])))
    return {
        "n_onduleurs": n_onduleurs,
        "puissance_necessaire_w": puissance_necessaire,
        "cout": n_onduleurs * onduleur_catalogue["prix"],
    }





# --- Calibration d'une chaîne de Markov depuis des données réelles --------

def estimer_chaine_markov_depuis_serie(serie, n_etats=3, matrice_defaut=None, etats_defaut=None,
                                        description="série"):
    """
    Calibre une chaîne de Markov à partir d'une série réelle : découpage en
    n_etats quantiles + comptage empirique des transitions observées.
    Repli automatique sur (etats_defaut, matrice_defaut) si la série est
    trop courte (< 3 x n_etats points) — jamais d'erreur bloquante.
    """
    seuil_minimal = 3 * n_etats
    if serie is None or len(serie) < seuil_minimal:
        n_recu = 0 if serie is None else len(serie)
        message = (
            f"Données insuffisantes pour calibrer la chaîne de {description} "
            f"({n_recu} valeur(s) reçue(s), {seuil_minimal} recommandées) : "
            f"utilisation de la matrice par défaut de la structure test."
        )
        return np.array(etats_defaut), np.array(matrice_defaut), message

    serie = np.asarray(serie, dtype=float)
    frontieres = np.quantile(serie, np.linspace(0, 1, n_etats + 1))
    frontieres[0] -= 1e-9
    indices_etats = np.clip(np.digitize(serie, frontieres) - 1, 0, n_etats - 1)

    etats = np.array([serie[indices_etats == k].mean() if np.any(indices_etats == k)
                       else serie.mean() for k in range(n_etats)])

    comptages = np.ones((n_etats, n_etats))
    for i in range(len(indices_etats) - 1):
        comptages[indices_etats[i], indices_etats[i + 1]] += 1
    matrice = comptages / comptages.sum(axis=1, keepdims=True)

    ordre = np.argsort(etats)
    etats = etats[ordre]
    matrice = matrice[np.ix_(ordre, ordre)]
    return etats, matrice, None


def stationary_distribution(matrice):
    """Distribution stationnaire pi d'une chaîne de Markov (pi.P = pi)."""
    matrice = np.asarray(matrice, dtype=float)
    valeurs_propres, vecteurs_propres = np.linalg.eig(matrice.T)
    idx = np.argmin(np.abs(valeurs_propres - 1))
    v = np.real(vecteurs_propres[:, idx])
    return v / v.sum()


def usure_panneaux(t_jours, lam_annuel=0.00893):
    """Facteur de dégradation des panneaux (80% de puissance garantie à 25 ans)."""
    return np.exp(-(lam_annuel / 365.0) * t_jours)


def _step_markov(etats_courants, matrice):
    """Fait avancer d'un pas un lot vectorisé de trajectoires de Markov."""
    n = etats_courants.shape[0]
    tirage = np.random.rand(n)
    cumule = np.cumsum(matrice[etats_courants], axis=1)
    return np.clip((tirage[:, None] > cumule).sum(axis=1), 0, matrice.shape[0] - 1)


# --- Simulation principale (Eval_consommation vectorisé) -------------------

def simuler_installation(n_panneaux, n_batteries, n_repetitions, parametres,
                          n_jours=365, lolp_cible=0.05, appliquer_degradation=False,
                          jours_par_etat_conso=30, soc_initial_frac=0.7, fenetre_lolp=365):
    """
    Simule n_repetitions trajectoires indépendantes (Monte-Carlo) d'une
    installation de n_panneaux panneaux et n_batteries batteries, sur
    n_jours jours (365 = une année ; un multiple de 365 pour évaluer une
    installation existante sur plusieurs années avec dégradation).

    IMPORTANT (fenêtre glissante) : le LOLP utilisé pour la condition
    d'arrêt est calculé sur une fenêtre GLISSANTE des `fenetre_lolp`
    derniers jours (365 par défaut), pas cumulé depuis le jour 0. Sur une
    simulation pluriannuelle avec dégradation des panneaux, un LOLP cumulé
    depuis le début dilue une dégradation récente dans la moyenne de
    plusieurs années et retarderait artificiellement la détection du
    moment où l'installation devient réellement insuffisante. La fenêtre
    glissante répond à la question "l'installation est-elle encore
    fiable EN CE MOMENT ?", pas "l'a-t-elle été en moyenne depuis sa mise
    en service ?".
    """
    p = parametres
    pc_installee = n_panneaux * p["pc_panneau_wc"]
    capacite_utile = n_batteries * p["capacite_batterie_ah"] * p["tension_bus_v"] * p["profondeur_decharge"]

    pi_cs = stationary_distribution(p["cs_matrice"])
    pi_psh = stationary_distribution(p["psh_matrice"])

    etat_cs = np.random.choice(len(p["cs_etats"]), size=n_repetitions, p=pi_cs)
    etat_psh = np.random.choice(len(p["psh_etats"]), size=n_repetitions, p=pi_psh)
    soc = np.full(n_repetitions, capacite_utile * soc_initial_frac)

    nb_jours_deficit = np.zeros(n_repetitions, dtype=int)          # cumul depuis le debut (indicatif)
    deficits_consecutifs = np.zeros(n_repetitions, dtype=int)
    jour_arret = np.full(n_repetitions, n_jours, dtype=int)
    declenche = np.zeros(n_repetitions, dtype=bool)
    somme_manque = np.zeros(n_repetitions)
    deficits_premiere_annee = np.zeros(n_repetitions, dtype=int)   # snapshot a t=365

    # Tampon circulaire pour la fenetre glissante du LOLP
    fenetre = min(fenetre_lolp, n_jours)
    tampon = np.zeros((fenetre, n_repetitions), dtype=int)
    somme_glissante = np.zeros(n_repetitions, dtype=int)

    cs_etats = np.asarray(p["cs_etats"])
    psh_etats = np.asarray(p["psh_etats"])
    cs_matrice = np.asarray(p["cs_matrice"])
    psh_matrice = np.asarray(p["psh_matrice"])

    for t in range(n_jours):
        if t % jours_par_etat_conso == 0:
            etat_cs = _step_markov(etat_cs, cs_matrice)
        etat_psh = _step_markov(etat_psh, psh_matrice)

        conso_jour = p["e_jour_base_wh"] * cs_etats[etat_cs]
        facteur_usure = usure_panneaux(t) if appliquer_degradation else 1.0
        prod_jour = pc_installee * psh_etats[etat_psh] * p["eta_systeme"] * facteur_usure

        solde = soc + prod_jour - conso_jour
        manque = np.clip(-solde, 0, None)
        soc = np.clip(solde, 0, capacite_utile)

        en_deficit = (manque > 0).astype(int)
        nb_jours_deficit += en_deficit
        deficits_consecutifs = np.where(en_deficit, deficits_consecutifs + 1, 0)
        somme_manque += manque

        idx = t % fenetre
        somme_glissante += en_deficit - tampon[idx]
        tampon[idx] = en_deficit

        if t == min(364, n_jours - 1):
            deficits_premiere_annee = nb_jours_deficit.copy()

        lolp_glissant = somme_glissante / min(t + 1, fenetre)
        nouveau = (~declenche) & ((((t + 1) >= fenetre) & (lolp_glissant > lolp_cible)) | (deficits_consecutifs >= 5))
        jour_arret[nouveau] = t + 1
        declenche |= nouveau

    fiabilite = 1.0 - nb_jours_deficit / n_jours
    fiabilite_premiere_annee = 1.0 - deficits_premiere_annee / min(365, n_jours)
    puissance_manquante = np.divide(somme_manque, nb_jours_deficit,
                                     out=np.zeros_like(somme_manque), where=nb_jours_deficit > 0)
    return {
        "fiabilite": fiabilite,
        "fiabilite_premiere_annee": fiabilite_premiere_annee,
        "puissance_manquante": puissance_manquante,
        "jour_arret": jour_arret,
        "nb_jours_deficit": nb_jours_deficit,
        "declenche": declenche,
    }


# --- Construction des paramètres (fusion données utilisateur + défauts) ---

def construire_parametres(fichier_consommation=None, fichier_meteo=None, avertissements=None):
    """
    Construit le dictionnaire de paramètres à utiliser pour la simulation :
    donnée fournie -> utilisée ; donnée absente/invalide -> repli automatique
    sur la structure test, jamais d'erreur bloquante.
    """
    if avertissements is None:
        avertissements = []

    parametres = {
        "e_jour_base_wh": DEFAUTS["e_jour_base_wh"],
        "eta_systeme": DEFAUTS["eta_systeme"],
        "pc_panneau_wc": DEFAUTS["pc_panneau_wc"],
        "capacite_batterie_ah": DEFAUTS["capacite_batterie_ah"],
        "tension_bus_v": DEFAUTS["tension_bus_v"],
        "profondeur_decharge": DEFAUTS["profondeur_decharge"],
    }

    if fichier_consommation:
        try:
            serie = charger_consommation_mensuelle(fichier_consommation)
            etats, matrice, msg = estimer_chaine_markov_depuis_serie(
                serie, n_etats=3, matrice_defaut=DEFAUTS["cs_matrice"],
                etats_defaut=DEFAUTS["cs_etats"], description="consommation"
            )
            moyenne = float(np.mean(etats))
            parametres["cs_etats"] = (np.array(etats) / moyenne).tolist()
            parametres["cs_matrice"] = matrice.tolist()
            parametres["e_jour_base_wh"] = moyenne / 30.0
            if msg:
                avertissements.append(msg)
        except DualSolarStatError as e:
            avertissements.append(f"{e} -> valeurs de consommation par défaut utilisées.")
            parametres["cs_etats"] = DEFAUTS["cs_etats"]
            parametres["cs_matrice"] = DEFAUTS["cs_matrice"]
    else:
        avertissements.append("Aucun fichier de consommation fourni : matrice de consommation par défaut utilisée.")
        parametres["cs_etats"] = DEFAUTS["cs_etats"]
        parametres["cs_matrice"] = DEFAUTS["cs_matrice"]

    if fichier_meteo:
        try:
            serie = charger_ensoleillement_journalier(fichier_meteo)
            etats, matrice, msg = estimer_chaine_markov_depuis_serie(
                serie, n_etats=3, matrice_defaut=DEFAUTS["psh_matrice"],
                etats_defaut=DEFAUTS["psh_etats"], description="ensoleillement"
            )
            parametres["psh_etats"] = np.array(etats).tolist()
            parametres["psh_matrice"] = matrice.tolist()
            if msg:
                avertissements.append(msg)
        except DualSolarStatError as e:
            avertissements.append(f"{e} -> valeurs météo par défaut utilisées.")
            parametres["psh_etats"] = DEFAUTS["psh_etats"]
            parametres["psh_matrice"] = DEFAUTS["psh_matrice"]
    else:
        avertissements.append("Aucun fichier météo fourni : matrice météo par défaut utilisée.")
        parametres["psh_etats"] = DEFAUTS["psh_etats"]
        parametres["psh_matrice"] = DEFAUTS["psh_matrice"]

    return parametres, avertissements


def resoudre_type_menage(type_menage):
    """Renvoie le LOLP cible pour un type de ménage, avec repli sur 'moyenne' si invalide."""
    if not type_menage:
        return TYPE_MENAGE_PAR_DEFAUT, TYPES_MENAGE[TYPE_MENAGE_PAR_DEFAUT]["lolp_cible"], None
    cle = str(type_menage).strip().lower()
    if cle not in TYPES_MENAGE:
        avertissement = (
            f"Type de ménage « {type_menage} » non reconnu (valeurs possibles : "
            f"{list(TYPES_MENAGE.keys())}) : type « {TYPE_MENAGE_PAR_DEFAUT} » utilisé par défaut."
        )
        return TYPE_MENAGE_PAR_DEFAUT, TYPES_MENAGE[TYPE_MENAGE_PAR_DEFAUT]["lolp_cible"], avertissement
    return cle, TYPES_MENAGE[cle]["lolp_cible"], None


# --- Fonctions de haut niveau (appelées depuis l'interface graphique) -----

def executer_mode_test(n_repetitions=1500):
    """MODE TEST / DÉMONSTRATION — reproduit le cas Mairie de Boulsa avec les valeurs par défaut."""
    return executer_evaluation({
        "n_panneaux": DEFAUTS["n_panneaux_test"],
        "n_batteries": DEFAUTS["n_batteries_test"],
        "fichier_consommation": None,
        "fichier_meteo": None,
        "type_menage": TYPE_MENAGE_PAR_DEFAUT,
        "n_repetitions": n_repetitions,
    })


def estimer_renforcement_necessaire(n_panneaux_actuel, n_batteries_actuel, parametres, objectif,
                                     n_sim=300, plafond_multiplicateur=2.5, surface_max_m2=None):
    """
    Quand une installation évaluée est insuffisante, cherche la meilleure
    combinaison de renforcement en faisant varier à la fois le nombre de
    panneaux et le nombre de batteries, sans figer l'un ou l'autre.

    La recherche est ordonnée par accroissement minimal du nombre de
    composants, ce qui privilégie les solutions les plus proches de la
    configuration actuelle.
    """
    surface_actuelle = estimer_surface_installation(n_panneaux_actuel, parametres["pc_panneau_wc"])
    plafond_panneaux = int(n_panneaux_actuel * plafond_multiplicateur) + 10
    plafond_batteries = int(n_batteries_actuel * plafond_multiplicateur) + 10
    meilleure_config_objectif = None
    meilleur_cout_objectif = None
    meilleure_config_possible = None
    meilleure_fiabilite_possible = -1.0

    for total_delta in range(1, plafond_panneaux - n_panneaux_actuel + plafond_batteries - n_batteries_actuel + 1):
        for delta_panneaux in range(0, min(total_delta, plafond_panneaux - n_panneaux_actuel) + 1):
            delta_batteries = total_delta - delta_panneaux
            if delta_batteries > plafond_batteries - n_batteries_actuel:
                continue

            n_pan = n_panneaux_actuel + delta_panneaux
            n_bat = n_batteries_actuel + delta_batteries
            surface_totale = estimer_surface_installation(n_pan, parametres["pc_panneau_wc"])
            surface_additionnelle = surface_totale - surface_actuelle
            if surface_max_m2 is not None and surface_additionnelle > surface_max_m2:
                continue

            r = simuler_installation(n_pan, n_bat, n_sim, parametres, lolp_cible=1 - objectif)
            fiabilite = float(r["fiabilite"].mean())

            if fiabilite > meilleure_fiabilite_possible or (
                fiabilite == meilleure_fiabilite_possible and total_delta < (meilleure_config_possible["panneaux_total_necessaire"] - n_panneaux_actuel + meilleure_config_possible["batteries_total_necessaire"] - n_batteries_actuel if meilleure_config_possible else float('inf'))
            ):
                meilleure_config_possible = {
                    "panneaux_total_necessaire": n_pan,
                    "panneaux_supplementaires": delta_panneaux,
                    "batteries_total_necessaire": n_bat,
                    "batteries_supplementaires": delta_batteries,
                    "surface_m2": surface_totale,
                    "surface_additionnelle_m2": surface_additionnelle,
                    "fiabilite_estimee": fiabilite,
                    "atteint_objectif": False,
                }
                meilleure_fiabilite_possible = fiabilite

            if fiabilite >= objectif:
                if meilleure_config_objectif is None or total_delta < meilleur_cout_objectif:
                    meilleure_config_objectif = {
                        "panneaux_total_necessaire": n_pan,
                        "panneaux_supplementaires": delta_panneaux,
                        "batteries_total_necessaire": n_bat,
                        "batteries_supplementaires": delta_batteries,
                        "surface_m2": surface_totale,
                        "surface_additionnelle_m2": surface_additionnelle,
                        "fiabilite_estimee": fiabilite,
                        "atteint_objectif": True,
                    }
                    meilleur_cout_objectif = total_delta
        if meilleure_config_objectif is not None:
            break

    if meilleure_config_objectif is not None:
        return meilleure_config_objectif

    return meilleure_config_possible if meilleure_config_possible is not None else {
        "panneaux_total_necessaire": None,
        "panneaux_supplementaires": None,
        "batteries_total_necessaire": None,
        "batteries_supplementaires": None,
        "surface_m2": surface_actuelle,
        "surface_additionnelle_m2": 0.0,
        "fiabilite_estimee": 0.0,
        "atteint_objectif": False,
    }


def executer_evaluation(config):
    """
    MODE ÉVALUATION D'UNE INSTALLATION EXISTANTE. Ne lève jamais d'exception :
    toute erreur est renvoyée dans le champ "succes"/"erreur" du dict retourné.

    Simule sur plusieurs années (config["duree_annees"], 15 ans par défaut)
    AVEC dégradation des panneaux, pour répondre à la question centrale du
    mode Eval_consommation : non seulement "cette installation couvre-t-elle
    le besoin aujourd'hui ?" mais aussi "dans combien de temps cessera-t-elle
    de le faire, à mesure que les panneaux vieillissent ?". Le LOLP utilisé
    pour détecter ce moment est calculé sur une fenêtre glissante d'un an
    (pas cumulé depuis la mise en service), afin que le diagnostic reflète
    l'état RÉCENT de l'installation plutôt qu'une moyenne historique diluée.
    """
    avertissements = []
    try:
        n_panneaux = int(config["n_panneaux"])
        n_batteries = int(config["n_batteries"])
        if n_panneaux <= 0 or n_batteries <= 0:
            raise DualSolarStatError("Le nombre de panneaux et de batteries doit être strictement positif.")
    except (KeyError, ValueError, TypeError):
        return {"succes": False, "erreur": "Nombre de panneaux et/ou de batteries manquant ou invalide."}

    try:
        duree_annees = float(config.get("duree_annees", 15))
        if duree_annees <= 0:
            raise ValueError
    except (ValueError, TypeError):
        return {"succes": False, "erreur": "La durée d'horizon (en années) doit être un nombre positif."}

    type_menage_utilise, lolp_cible, avert_menage = resoudre_type_menage(config.get("type_menage"))
    if avert_menage:
        avertissements.append(avert_menage)

    try:
        parametres, avertissements = construire_parametres(
            fichier_consommation=config.get("fichier_consommation"),
            fichier_meteo=config.get("fichier_meteo"),
            avertissements=avertissements,
        )
    except Exception as e:
        return {"succes": False, "erreur": f"Erreur inattendue lors de la préparation des données : {e}"}

    surface_max_m2 = config.get("surface_max_m2")
    if surface_max_m2 is not None:
        try:
            surface_max_m2 = float(surface_max_m2)
            if surface_max_m2 <= 0:
                raise ValueError
        except (TypeError, ValueError):
            return {"succes": False, "erreur": "La surface maximale doit être un nombre positif."}

    n_repetitions = int(config.get("n_repetitions", 800))
    objectif = 1 - lolp_cible

    # Phase 1 : vérification rapide sur une seule année (dégradation
    # négligeable la première année), pour savoir si l'installation couvre
    # DEJA le besoin aujourd'hui, avant de lancer une simulation longue.
    resultats_annee_1 = simuler_installation(n_panneaux, n_batteries, n_repetitions, parametres,
                                              n_jours=365, lolp_cible=lolp_cible, appliquer_degradation=False)
    fiabilite_initiale = float(resultats_annee_1["fiabilite"].mean())
    suffisant_au_depart = fiabilite_initiale >= objectif

    renforcement = None
    annees_avant_insuffisance = None
    part_qui_degradent = None
    puissance_manquante_moyenne_wh = float(
        resultats_annee_1["puissance_manquante"][resultats_annee_1["nb_jours_deficit"] > 0].mean()
    ) if np.any(resultats_annee_1["nb_jours_deficit"] > 0) else 0.0

    if not suffisant_au_depart:
        # Sous-dimensionnement actuel, dès la première année : ce n'est pas
        # un problème de dégradation future. On calcule le renforcement
        # immédiat nécessaire (snapshot à l'état neuf) et on cherche une
        # combinaison conjointe de panneaux et batteries.
        renforcement = estimer_renforcement_necessaire(
            n_panneaux, n_batteries, parametres, objectif,
            surface_max_m2=surface_max_m2,
        )
    else:
        # Phase 2 : l'installation est suffisante aujourd'hui. On simule sur
        # tout l'horizon (duree_annees, dégradation des panneaux activée)
        # pour savoir au bout de combien de temps elle cesse de l'être.
        n_jours = max(365, int(round(duree_annees * 365)))
        resultats_long_terme = simuler_installation(n_panneaux, n_batteries, n_repetitions, parametres,
                                                      n_jours=n_jours, lolp_cible=lolp_cible,
                                                      appliquer_degradation=True)
        jours_declenches = resultats_long_terme["jour_arret"][resultats_long_terme["declenche"]]
        part_qui_degradent = float(resultats_long_terme["declenche"].mean())
        if len(jours_declenches):
            annees_avant_insuffisance = float(jours_declenches.mean()) / 365.0

    return {
        "succes": True,
        "n_panneaux": n_panneaux,
        "n_batteries": n_batteries,
        "type_menage_utilise": type_menage_utilise,
        "lolp_cible": lolp_cible,
        "duree_annees_simulee": duree_annees,
        "fiabilite_initiale": fiabilite_initiale,
        "fiabilite_ecart_type": float(resultats_annee_1["fiabilite"].std()),
        "suffisant_au_depart": suffisant_au_depart,
        "puissance_manquante_moyenne_wh": puissance_manquante_moyenne_wh,
        "part_qui_degradent_sous_objectif": part_qui_degradent,
        "annees_avant_insuffisance": annees_avant_insuffisance,
        "renforcement": renforcement,
        "pc_panneau_wc": parametres["pc_panneau_wc"],
        "avertissements": avertissements,
    }


def executer_dimensionnement(config):
    """
    MODE DIMENSIONNEMENT D'UNE NOUVELLE INSTALLATION. Nécessite un catalogue
    matériel CSV (config["fichier_catalogue"]).
    """
    if not config.get("fichier_catalogue"):
        return {"succes": False, "erreur": "Le catalogue matériel (CSV) est obligatoire pour le dimensionnement."}

    try:
        catalogue = charger_catalogue_materiel(config["fichier_catalogue"])
    except DualSolarStatError as e:
        return {"succes": False, "erreur": str(e)}

    try:
        budget_max = float(config["budget_max"])
        if budget_max <= 0:
            raise ValueError
    except (KeyError, ValueError, TypeError):
        return {"succes": False, "erreur": "Budget maximal manquant ou invalide."}

    type_menage_utilise, lolp_cible, avert_menage = resoudre_type_menage(config.get("type_menage"))
    avertissements = [avert_menage] if avert_menage else []

    parametres, avertissements = construire_parametres(
        fichier_consommation=config.get("fichier_consommation"),
        fichier_meteo=config.get("fichier_meteo"),
        avertissements=avertissements,
    )

    surface_max_m2 = config.get("surface_max_m2")
    if surface_max_m2 is not None and surface_max_m2 != "":
        surface_max_m2 = float(surface_max_m2)
    else:
        surface_max_m2 = None

    n_sim = int(config.get("n_repetitions_par_config", 300))
    objectif = 1.0 - lolp_cible

    meilleure_config, meilleur_cout = None, None
    for panneau in catalogue["panneaux"]:
        for batterie in catalogue["batteries"]:
            for onduleur in catalogue["onduleurs"]:
                parametres_loc = dict(parametres)
                parametres_loc["pc_panneau_wc"] = panneau["caracteristique"]
                parametres_loc["capacite_batterie_ah"] = batterie["caracteristique"]

                for n_pan in range(4, 61):
                    surface = estimer_surface_installation(n_pan, panneau["caracteristique"])
                    if surface_max_m2 is not None and surface > surface_max_m2:
                        break

                    cout_panneaux = n_pan * panneau["prix"]
                    if cout_panneaux > budget_max:
                        break

                    info_onduleur = dimensionner_onduleur(n_pan * panneau["caracteristique"], onduleur)
                    cout_panneaux_et_onduleurs = cout_panneaux + info_onduleur["cout"]
                    if cout_panneaux_et_onduleurs > budget_max:
                        continue

                    for n_bat in range(1, 21):
                        cout_total = cout_panneaux_et_onduleurs + n_bat * batterie["prix"]
                        if cout_total > budget_max:
                            break

                        resultats = simuler_installation(n_pan, n_bat, n_sim, parametres_loc, lolp_cible=lolp_cible)
                        fiab = float(resultats["fiabilite"].mean())
                        if fiab >= objectif and (meilleur_cout is None or cout_total < meilleur_cout):
                            meilleur_cout = cout_total
                            meilleure_config = {
                                "n_panneaux": n_pan,
                                "n_batteries": n_bat,
                                "fiabilite": fiab,
                                "cout": cout_total,
                                "modele_panneau": panneau["nom"],
                                "modele_batterie": batterie["nom"],
                                "n_onduleurs": info_onduleur["n_onduleurs"],
                                "modele_onduleur": onduleur["nom"],
                                "cout_panneaux": cout_panneaux,
                                "cout_batteries": n_bat * batterie["prix"],
                                "cout_onduleurs": info_onduleur["cout"],
                                "puissance_onduleur_necessaire_w": info_onduleur["puissance_necessaire_w"],
                                "surface_estimee_m2": surface,
                            }

    if meilleure_config is None:
        return {
            "succes": False,
            "erreur": (f"Aucune configuration ne permet d'atteindre la fiabilité cible "
                       f"({objectif:.0%}) avec un budget de {budget_max:,.0f}. "
                       f"Augmentez le budget, la surface disponible ou choisissez un type de ménage moins exigeant."),
            "avertissements": avertissements,
        }

    return {"succes": True, "configuration_recommandee": meilleure_config,
            "type_menage_utilise": type_menage_utilise, "lolp_cible": lolp_cible,
            "avertissements": avertissements}

    return {"succes": True, "configuration_recommandee": meilleure_config,
            "type_menage_utilise": type_menage_utilise, "lolp_cible": lolp_cible,
            "avertissements": avertissements}


def generer_fichiers_exemple(dossier="."):
    """Écrit 3 fichiers CSV d'exemple (bon format), remplis de données plausibles."""
    with open(os.path.join(dossier, "exemple_consommation_mensuelle.csv"), "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["mois", "consommation_wh"])
        cs_moyen_reference = sum(DEFAUTS["cs_etats"]) / len(DEFAUTS["cs_etats"])
        base_mensuelle = DEFAUTS["e_jour_base_wh"] * cs_moyen_reference * 30
        facteurs_saisonniers = [0.92, 0.95, 1.05, 1.00, 0.90, 1.10,
                                 1.18, 1.16, 1.02, 0.97, 0.93, 0.90]
        for i, facteur in enumerate(facteurs_saisonniers, start=1):
            w.writerow([f"2025-{i:02d}", round(base_mensuelle * facteur)])

    with open(os.path.join(dossier, "exemple_ensoleillement_journalier.csv"), "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["jour", "psh_heures"])
        rng = np.random.default_rng(0)
        for i, v in enumerate(np.clip(rng.normal(5.3, 1.2, 60), 1.0, 7.5), start=1):
            w.writerow([f"jour_{i}", round(float(v), 2)])

    with open(os.path.join(dossier, "exemple_catalogue_materiel.csv"), "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["type", "nom", "caracteristique", "prix"])
        w.writerow(["panneau", "Panneau monocristallin 230Wc", 230, 77000])
        w.writerow(["panneau", "Panneau monocristallin 400Wc", 400, 220000])
        w.writerow(["batterie", "Batterie Lithium LiFePO4 51.2V 100Ah", 100, 570000])
        w.writerow(["batterie", "Batterie Lithium LiFePO4 51.2V 400Ah", 400, 5400000])
        w.writerow(["batterie", "Batterie Lithium LiFePO4 51.2V 200Ah", 200, 1800000])
        w.writerow(["onduleur", "Onduleur hybride 7kW MPPT intégré", 7000, 945000])
        w.writerow(["onduleur", "Onduleur hybride 5kW MPPT intégré", 5000, 820000])
        w.writerow(["onduleur", "Onduleur hybride 3kW MPPT intégré", 3000, 620000])


# ============================================================================
# ==== SECTION B : INTERFACE GRAPHIQUE (CustomTkinter) ======================
# ============================================================================

ctk.set_appearance_mode("light")       # "system", "light" ou "dark" (clair par défaut)
ctk.set_default_color_theme("green")    # theme visuel (vert = coherent avec le solaire)

COULEUR_AVERTISSEMENT = "#D9A441"
COULEUR_ERREUR = "#E5484D"
COULEUR_TITRE = "#2FA572"
COULEUR_TEXTE = ("#1A1A1A", "#DCE4E2")  # (mode clair, mode sombre)


class SelecteurFichier(ctk.CTkFrame):
    """Ligne réutilisable : label + chemin du fichier choisi + boutons Parcourir/Effacer."""

    def __init__(self, parent, libelle, obligatoire=False, **kwargs):
        super().__init__(parent, fg_color="transparent", **kwargs)
        self.chemin = ctk.StringVar(value="")
        texte_libelle = libelle + (" *" if obligatoire else " (optionnel)")

        self.grid_columnconfigure(1, weight=1)
        ctk.CTkLabel(self, text=texte_libelle, width=230, anchor="w").grid(row=0, column=0, sticky="w")
        self.champ = ctk.CTkEntry(self, textvariable=self.chemin, placeholder_text="Aucun fichier sélectionné")
        self.champ.grid(row=0, column=1, sticky="ew", padx=6)
        ctk.CTkButton(self, text="Parcourir...", width=100, command=self._parcourir).grid(row=0, column=2, padx=2)
        ctk.CTkButton(self, text="✕", width=32, fg_color="transparent", border_width=1,
                      command=self._effacer).grid(row=0, column=3, padx=2)

    def _parcourir(self):
        chemin = filedialog.askopenfilename(
            title="Choisir un fichier CSV",
            filetypes=[("Fichiers CSV", "*.csv"), ("Tous les fichiers", "*.*")],
        )
        if chemin:
            self.chemin.set(chemin)

    def _effacer(self):
        self.chemin.set("")

    def valeur(self):
        v = self.chemin.get().strip()
        return v if v else None


class FenetreResultats(ctk.CTkToplevel):
    """Fenêtre dédiée pour afficher les résultats avec beaucoup d'espace."""

    def __init__(self, parent, contenu="", resultat_dict=None, titre="Résultats"):
        super().__init__(parent)
        self.title(titre)
        self.geometry("900x650")
        self.minsize(780, 520)
        self.transient(parent)

        self.dernier_resultat = resultat_dict
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)

        entete = ctk.CTkFrame(self, fg_color="transparent")
        entete.grid(row=0, column=0, sticky="ew", padx=12, pady=(10, 6))
        entete.grid_columnconfigure(0, weight=1)
        ctk.CTkLabel(entete, text="Résultats détaillés", font=ctk.CTkFont(size=14, weight="bold")).grid(row=0, column=0, sticky="w")
        self.bouton_export = ctk.CTkButton(entete, text="💾 Exporter...", width=120,
                                            command=self._exporter, state="disabled")
        self.bouton_export.grid(row=0, column=1, sticky="e")

        self.texte = ctk.CTkTextbox(self, wrap="word", font=("Consolas", 12), activate_scrollbars=True)
        self.texte.grid(row=1, column=0, sticky="nsew", padx=12, pady=(0, 12))
        self.texte.configure(state="disabled")

        self.texte.tag_config("titre", foreground=COULEUR_TITRE)
        self.texte.tag_config("avertissement", foreground=COULEUR_AVERTISSEMENT)
        self.texte.tag_config("erreur", foreground=COULEUR_ERREUR)

        self.actualiser(contenu, resultat_dict)

    def actualiser(self, contenu="", resultat_dict=None):
        self.dernier_resultat = resultat_dict
        self.texte.configure(state="normal")
        self.texte.delete("1.0", "end")
        if contenu:
            self.texte.insert("end", contenu)
        self.texte.configure(state="disabled")
        self.texte.see("end")
        self.bouton_export.configure(state="normal" if resultat_dict is not None else "disabled")

    def _exporter(self):
        exporter_resultat(self.dernier_resultat, self.texte.get("1.0", "end"))


class ZoneResultats(ctk.CTkFrame):
    """Zone de résultats : titre + texte formaté + bouton d'export."""

    def __init__(self, parent, **kwargs):
        super().__init__(parent, **kwargs)
        self.dernier_resultat = None  # dict brut, pour l'export CSV
        self.fenetre_popup = None
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(1, weight=1)

        entete = ctk.CTkFrame(self, fg_color="transparent")
        entete.grid(row=0, column=0, sticky="ew", padx=10, pady=(10, 4))
        entete.grid_columnconfigure(0, weight=1)
        ctk.CTkLabel(entete, text="Résultats", font=ctk.CTkFont(size=14, weight="bold")).grid(row=0, column=0, sticky="w")
        self.bouton_export = ctk.CTkButton(entete, text="💾 Exporter...", width=120,
                                            command=self._exporter, state="disabled")
        self.bouton_export.grid(row=0, column=1, sticky="e")
        self.bouton_fenetre = ctk.CTkButton(entete, text="🗗 Nouvelle fenêtre", width=140,
                                             command=self._ouvrir_dans_nouvelle_fenetre)
        self.bouton_fenetre.grid(row=0, column=2, sticky="e", padx=(8, 0))

        self.texte = ctk.CTkTextbox(self, wrap="word", font=("Consolas", 12), activate_scrollbars=True)
        self.texte.grid(row=1, column=0, sticky="nsew", padx=10, pady=(0, 10))
        self.texte.configure(state="disabled")

        self.texte.tag_config("titre", foreground=COULEUR_TITRE)
        self.texte.tag_config("avertissement", foreground=COULEUR_AVERTISSEMENT)
        self.texte.tag_config("erreur", foreground=COULEUR_ERREUR)

    def effacer(self):
        self.texte.configure(state="normal")
        self.texte.delete("1.0", "end")
        self.texte.configure(state="disabled")
        self.dernier_resultat = None
        self.bouton_export.configure(state="disabled")

    def ecrire(self, contenu, tag=None):
        self.texte.configure(state="normal")
        if tag:
            self.texte.insert("end", contenu + "\n", tag)
        else:
            self.texte.insert("end", contenu + "\n")
        self.texte.configure(state="disabled")
        self.texte.see("end")

    def definir_resultat_exportable(self, resultat_dict):
        """Mémorise le dernier résultat structuré (dict) pour permettre l'export CSV."""
        self.dernier_resultat = resultat_dict
        self.bouton_export.configure(state="normal")

    def contenu(self):
        return self.texte.get("1.0", "end")

    def _ouvrir_dans_nouvelle_fenetre(self):
        contenu = self.contenu()
        if not contenu.strip():
            return
        if self.fenetre_popup is not None and self.fenetre_popup.winfo_exists():
            self.fenetre_popup.actualiser(contenu, self.dernier_resultat)
            self.fenetre_popup.focus()
        else:
            self.fenetre_popup = FenetreResultats(self.winfo_toplevel(), contenu, self.dernier_resultat,
                                                 titre="Résultats détaillés")
            self.fenetre_popup.focus()

    def _exporter(self):
        exporter_resultat(self.dernier_resultat, self.texte.get("1.0", "end"))


# --- Export des résultats (.txt ou .csv) -----------------------------------

def _aplatir_dict(d, prefixe=""):
    """Aplatit un dict (éventuellement imbriqué / avec listes) en paires clé/valeur pour un CSV."""
    lignes = []
    for cle, valeur in d.items():
        nom = f"{prefixe}{cle}"
        if isinstance(valeur, dict):
            lignes.extend(_aplatir_dict(valeur, prefixe=nom + "."))
        elif isinstance(valeur, (list, tuple)):
            lignes.append((nom, "; ".join(str(v) for v in valeur)))
        else:
            lignes.append((nom, valeur))
    return lignes


def _generer_docx(chemin, contenu_texte, resultat_dict=None, image_path=None):
    """Crée un document Word avec le texte, un commentaire de schéma, et une image éventuelle."""
    document = Document()
    document.add_heading("Rapport DualSolarStat", level=1)

    if contenu_texte:
        for ligne in contenu_texte.splitlines():
            if ligne.strip() == "":
                document.add_paragraph()
            else:
                document.add_paragraph(ligne)

    document.add_heading("Commentaires sur le schéma", level=2)
    document.add_paragraph(
        "Cette page contient une analyse de la simulation et du schéma de fiabilité. "
        "Le graphique inclus illustre la robustesse des recommandations de dimensionnement "
        "lorsque l'algorithme est relancé plusieurs fois avec la même configuration de référence."
    )

    if resultat_dict is not None:
        if resultat_dict.get("suffisant_au_depart") is True:
            document.add_paragraph(
                "L'évaluation montre que l'installation actuelle est suffisante à court terme, "
                "mais la dégradation des panneaux pourrait réduire cette fiabilité au fil des ans."
            )
        elif resultat_dict.get("suffisant_au_depart") is False:
            document.add_paragraph(
                "L'évaluation montre que l'installation actuelle n'est pas suffisant dès la première année. "
                "Un renforcement est nécessaire pour atteindre la fiabilité cible."
            )
        document.add_paragraph(
            "Le modèle repose sur des chaînes de Markov pour représenter les états de consommation "
            "et d'ensoleillement. Le schéma traduit la probabilité qu'un site reste fiable en fonction "
            "du temps et de l'incertitude climatique."
        )

    if image_path and os.path.isfile(image_path):
        try:
            document.add_heading("Figure générée", level=2)
            document.add_picture(image_path, width=Inches(6))
            document.add_paragraph(
                "La figure ci-dessus présente un boxplot des recommandations de dimensionnement obtenues "
                "sur plusieurs exécutions indépendantes. Une forte concentration des nombres de panneaux et "
                "de batteries autour d'une même valeur indique une recommandation robuste face au bruit de simulation."
            )
        except Exception:
            document.add_paragraph("[L'image n'a pas pu être insérée dans le document Word.]")

    document.save(chemin)


def exporter_resultat(resultat_dict, contenu_texte, image_path=None):
    """
    Ouvre une boîte de dialogue "Enregistrer sous..." et exporte soit le
    texte formaté (.txt), soit les résultats structurés (.csv), selon
    l'extension choisie par l'utilisateur.
    """
    chemin = filedialog.asksaveasfilename(
        title="Exporter les résultats",
        defaultextension=".txt",
        filetypes=[("Fichier texte", "*.txt"), ("Fichier CSV", "*.csv"), ("Document Word", "*.docx")],
    )
    if not chemin:
        return
    try:
        if chemin.lower().endswith(".csv"):
            if not resultat_dict:
                raise DualSolarStatError("Aucun résultat structuré à exporter en CSV pour cet onglet.")
            lignes = _aplatir_dict(resultat_dict)
            with open(chemin, "w", newline="", encoding="utf-8") as f:
                w = csv.writer(f)
                w.writerow(["champ", "valeur"])
                w.writerows(lignes)
        elif chemin.lower().endswith(".docx"):
            _generer_docx(chemin, contenu_texte, resultat_dict=resultat_dict, image_path=image_path)
        else:
            with open(chemin, "w", encoding="utf-8") as f:
                f.write(contenu_texte)
        messagebox.showinfo("Export réussi", f"Résultats exportés vers :\n{chemin}")
    except Exception as e:
        messagebox.showerror("Échec de l'export", f"Impossible d'exporter les résultats :\n{e}")


# --- Exécution en arrière-plan (thread-safe avec une file d'attente) ------

def lancer_en_arriere_plan(fenetre, bouton, barre_progression, fonction_calcul, fonction_affichage):
    """
    Exécute fonction_calcul() dans un thread séparé pour ne pas figer
    l'interface, puis affiche le résultat. IMPORTANT : le thread
    d'arrière-plan ne touche jamais un widget Tk directement (Tkinter
    n'est pas thread-safe) ; il dépose seulement son résultat dans une
    file d'attente que le thread principal sonde via .after().
    """
    bouton.configure(state="disabled")
    barre_progression.configure(mode="indeterminate")
    barre_progression.start()
    file_resultat = queue.Queue()

    def travail():
        try:
            resultat = fonction_calcul()
            file_resultat.put(("ok", resultat))
        except Exception as e:
            file_resultat.put(("erreur", str(e)))

    def sonder_file():
        try:
            statut, valeur = file_resultat.get_nowait()
        except queue.Empty:
            fenetre.after(100, sonder_file)
            return
        barre_progression.stop()
        bouton.configure(state="normal")
        if statut == "erreur":
            messagebox.showerror("Erreur inattendue", valeur)
        else:
            fonction_affichage(valeur)

    threading.Thread(target=travail, daemon=True).start()
    fenetre.after(100, sonder_file)


# --- Affichage des résultats (commun aux onglets) --------------------------

def afficher_resultat_evaluation(zone, resultat):
    zone.effacer()
    if not resultat.get("succes"):
        zone.ecrire("ÉCHEC DU CALCUL", "erreur")
        zone.ecrire(resultat.get("erreur", "Erreur inconnue."), "erreur")
        return

    zone.ecrire("=== ÉVALUATION DE L'INSTALLATION ===", "titre")
    zone.ecrire(f"Configuration testée : {resultat['n_panneaux']} panneaux / {resultat['n_batteries']} batteries")
    zone.ecrire(f"Type de ménage retenu : {resultat['type_menage_utilise']} "
                f"(LOLP cible = {resultat['lolp_cible']:.0%}, fiabilité cible = {1 - resultat['lolp_cible']:.0%})")
    zone.ecrire("")
    zone.ecrire(f"Fiabilité actuelle (1ère année, sans usure) : {resultat['fiabilite_initiale']:.1%} "
                f"(écart-type : {resultat['fiabilite_ecart_type']:.1%})")
    if resultat["puissance_manquante_moyenne_wh"] > 0:
        zone.ecrire(f"Énergie moyenne non couverte les jours de déficit : "
                    f"{resultat['puissance_manquante_moyenne_wh']:.0f} Wh")

    suffisant = resultat["suffisant_au_depart"]
    zone.ecrire("")
    if suffisant:
        zone.ecrire("Fiabilité actuelle suffisante ✓", "titre")
        zone.ecrire("")
        zone.ecrire(f"Projection sur {resultat['duree_annees_simulee']:.0f} ans (avec dégradation des panneaux, "
                    f"~0,89%/an) :", "titre")
        if resultat["annees_avant_insuffisance"] is not None:
            zone.ecrire(f" • L'installation cessera probablement de couvrir le besoin dans environ "
                        f"{resultat['annees_avant_insuffisance']:.1f} ans (durée à laquelle {resultat['part_qui_degradent_sous_objectif']:.0%} "
                        f"des trajectoires simulées passent sous l'objectif de fiabilité).", "avertissement")
            zone.ecrire(" • Prévoir un contrôle ou un renforcement de l'installation à cette échéance.")
        else:
            zone.ecrire(f" • Sur l'horizon simulé de {resultat['duree_annees_simulee']:.0f} ans, l'installation reste "
                        f"suffisante dans la quasi-totalité des trajectoires : pas de renforcement à prévoir "
                        f"à cet horizon.")
    else:
        zone.ecrire("Fiabilité actuelle INSUFFISANTE — renforcement recommandé dès maintenant ✗", "erreur")
        zone.ecrire("(l'installation ne couvre déjà pas le besoin cible dès la 1ère année : ce n'est pas "
                    "un problème de dégradation future, mais de dimensionnement actuel.)", "avertissement")

    if not suffisant and resultat.get("renforcement"):
        r = resultat["renforcement"]
        zone.ecrire("")
        if r.get("atteint_objectif"):
            zone.ecrire("Recommandation de renforcement conjointe (panneaux + batteries) :", "titre")
            zone.ecrire(
                f" • Passer à {r['panneaux_total_necessaire']} panneaux (+{r['panneaux_supplementaires']}) "
                f"et à {r['batteries_total_necessaire']} batteries (+{r['batteries_supplementaires']}), "
                "en respectant la surface maximale disponible si elle a été fournie."
            )
            zone.ecrire(f"   - Surface totale estimée nécessaire : {r['surface_m2']:.1f} m².")
            zone.ecrire(f"   - Surface additionnelle estimée nécessaire : {r['surface_additionnelle_m2']:.1f} m².")
            zone.ecrire(f"   - Fiabilité estimée atteinte : {r['fiabilite_estimee']:.1%}.")
            zone.ecrire("   - Objectif de fiabilité atteint sous les contraintes testées.")
        else:
            zone.ecrire("Proposition d'amélioration possible sous contrainte de surface :", "titre")
            if r["panneaux_total_necessaire"] is not None and r["batteries_total_necessaire"] is not None:
                zone.ecrire(
                    f" • Passer à {r['panneaux_total_necessaire']} panneaux (+{r['panneaux_supplementaires']}) "
                    f"et à {r['batteries_total_necessaire']} batteries (+{r['batteries_supplementaires']}), "
                    "en respectant la surface disponible aujourd'hui."
                )
                zone.ecrire(f"   - Surface totale estimée nécessaire : {r['surface_m2']:.1f} m².")
                zone.ecrire(f"   - Surface additionnelle estimée nécessaire : {r['surface_additionnelle_m2']:.1f} m².")
                zone.ecrire(f"   - Fiabilité maximale estimée avec cette surface : {r['fiabilite_estimee']:.1%}.")
                zone.ecrire(
                    "   - C'est la meilleure amélioration possible avec la surface actuelle, "
                    f"même si l'objectif de fiabilité cible ({1 - resultat['lolp_cible']:.0%}) n'est pas atteint."
                )
            else:
                surface_par_panneau = estimer_surface_installation(1, resultat.get('pc_panneau_wc', 600.0))
                zone.ecrire(
                    " • La surface additionnelle disponible n'est pas suffisante pour ajouter un panneau complet. "
                    f"Un panneau nécessite environ {surface_par_panneau:.1f} m² de surface."
                , "avertissement")
                zone.ecrire(
                    "   - Avec cet espace, l'ajout de batteries seules ne permettrait pas d'améliorer "
                    "significativement la fiabilité si la production PV reste la contrainte."
                , "avertissement")
        zone.ecrire(" • Pour un chiffrage précis (coût, modèles réels, onduleur inclus), utilisez "
                    "l'onglet « Dimensionner une nouvelle installation » avec votre catalogue.")

    if resultat["avertissements"]:
        zone.ecrire("")
        zone.ecrire("Avertissements :", "titre")
        for a in resultat["avertissements"]:
            zone.ecrire(" • " + a, "avertissement")

    zone.definir_resultat_exportable(resultat)
    zone._ouvrir_dans_nouvelle_fenetre()


def afficher_resultat_dimensionnement(zone, resultat):
    zone.effacer()
    if not resultat.get("succes"):
        zone.ecrire("ÉCHEC DU CALCUL", "erreur")
        zone.ecrire(resultat.get("erreur", "Erreur inconnue."), "erreur")
        if resultat.get("avertissements"):
            for a in resultat["avertissements"]:
                zone.ecrire(" • " + a, "avertissement")
        return

    c = resultat["configuration_recommandee"]
    zone.ecrire("=== DIMENSIONNEMENT RECOMMANDÉ ===", "titre")
    zone.ecrire(f"Type de ménage retenu : {resultat['type_menage_utilise']} "
                f"(LOLP cible = {resultat['lolp_cible']:.0%})")
    zone.ecrire("")
    zone.ecrire(f"Panneaux  : {c['n_panneaux']} x {c['modele_panneau']}  "
                f"({c['cout_panneaux']:,.0f})")
    zone.ecrire(f"Batteries : {c['n_batteries']} x {c['modele_batterie']}  "
                f"({c['cout_batteries']:,.0f})")
    zone.ecrire(f"Onduleurs : {c['n_onduleurs']} x {c['modele_onduleur']}  "
                f"({c['cout_onduleurs']:,.0f})")
    zone.ecrire(f"  (puissance onduleur nécessaire, marge de sécurité {MARGE_SECURITE_ONDULEUR:.0%} incluse : "
                f"{c['puissance_onduleur_necessaire_w']:.0f} W)")
    zone.ecrire("")
    zone.ecrire(f"Fiabilité obtenue : {c['fiabilite']:.1%}")
    zone.ecrire(f"Coût total estimé : {c['cout']:,.0f}")

    if resultat["avertissements"]:
        zone.ecrire("")
        zone.ecrire("Avertissements :", "titre")
        for a in resultat["avertissements"]:
            zone.ecrire(" • " + a, "avertissement")

    zone.definir_resultat_exportable(resultat)
    zone._ouvrir_dans_nouvelle_fenetre()


# --- Onglet 1 : Évaluation d'une installation existante ------------------

def construire_onglet_evaluation(onglet, fenetre_racine):
    onglet.grid_columnconfigure(0, weight=1)
    onglet.grid_rowconfigure(4, weight=1)

    cadre_saisie = ctk.CTkFrame(onglet)
    cadre_saisie.grid(row=0, column=0, sticky="ew", padx=15, pady=(15, 8))
    ctk.CTkLabel(cadre_saisie, text="Installation à évaluer", font=ctk.CTkFont(weight="bold")).grid(
        row=0, column=0, columnspan=2, sticky="w", padx=10, pady=(10, 6))

    ctk.CTkLabel(cadre_saisie, text="Nombre de panneaux installés (vide = 28 par défaut)").grid(row=1, column=0, sticky="w", padx=10, pady=4)
    champ_panneaux = ctk.CTkEntry(cadre_saisie, width=100)
    champ_panneaux.insert(0, "")
    champ_panneaux.grid(row=1, column=1, sticky="w", padx=10, pady=4)

    ctk.CTkLabel(cadre_saisie, text="Nombre de batteries installées (vide = 4 par défaut)").grid(row=2, column=0, sticky="w", padx=10, pady=4)
    champ_batteries = ctk.CTkEntry(cadre_saisie, width=100)
    champ_batteries.insert(0, "")
    champ_batteries.grid(row=2, column=1, sticky="w", padx=10, pady=4)

    ctk.CTkLabel(cadre_saisie, text="Surface maximale disponible pour l'amélioration (m²)").grid(row=3, column=0, sticky="w", padx=10, pady=4)
    champ_surface = ctk.CTkEntry(cadre_saisie, width=100)
    champ_surface.insert(0, "")
    champ_surface.grid(row=3, column=1, sticky="w", padx=10, pady=4)

    ctk.CTkLabel(cadre_saisie, text="Type de ménage / site").grid(row=4, column=0, sticky="w", padx=10, pady=4)
    valeurs_menage = [f"{cle} — {v['libelle']}" for cle, v in TYPES_MENAGE.items()]
    combo_menage = ctk.CTkComboBox(cadre_saisie, values=valeurs_menage, width=280, state="readonly")
    combo_menage.set(f"{TYPE_MENAGE_PAR_DEFAUT} — {TYPES_MENAGE[TYPE_MENAGE_PAR_DEFAUT]['libelle']}")
    combo_menage.grid(row=4, column=1, sticky="w", padx=10, pady=4)

    ctk.CTkLabel(cadre_saisie, text="Horizon de simulation (années, vide = 15 par défaut)").grid(row=4, column=0, sticky="w", padx=10, pady=(4, 10))
    champ_duree = ctk.CTkEntry(cadre_saisie, width=100)
    champ_duree.insert(0, "")
    champ_duree.grid(row=4, column=1, sticky="w", padx=10, pady=(4, 10))

    cadre_fichiers = ctk.CTkFrame(onglet)
    cadre_fichiers.grid(row=1, column=0, sticky="ew", padx=15, pady=8)
    ctk.CTkLabel(cadre_fichiers, text="Données réelles (optionnel — sinon valeurs du cas test)",
                 font=ctk.CTkFont(weight="bold")).pack(anchor="w", padx=10, pady=(10, 6))
    selecteur_conso = SelecteurFichier(cadre_fichiers, "Consommation mensuelle (CSV)")
    selecteur_conso.pack(fill="x", padx=10, pady=4)
    selecteur_meteo = SelecteurFichier(cadre_fichiers, "Ensoleillement journalier (CSV)")
    selecteur_meteo.pack(fill="x", padx=10, pady=(4, 10))

    barre = ctk.CTkProgressBar(onglet, mode="indeterminate")
    zone = ZoneResultats(onglet)

    def lancer():
        def valeur_ou_defaut(champ, defaut):
            texte = champ.get().strip()
            return defaut if texte == "" else texte

        try:
            n_panneaux = int(valeur_ou_defaut(champ_panneaux, DEFAUTS["n_panneaux_test"]))
            n_batteries = int(valeur_ou_defaut(champ_batteries, DEFAUTS["n_batteries_test"]))
        except ValueError:
            messagebox.showerror("Saisie invalide", "Le nombre de panneaux et de batteries doit être un entier.")
            return
        try:
            duree_annees = float(valeur_ou_defaut(champ_duree, 15))
            if duree_annees <= 0:
                raise ValueError
        except ValueError:
            messagebox.showerror("Saisie invalide", "L'horizon de simulation doit être un nombre d'années positif.")
            return

        type_menage = combo_menage.get().split(" — ")[0] if combo_menage.get() else TYPE_MENAGE_PAR_DEFAUT
        surface_max = champ_surface.get().strip()
        config = {
            "n_panneaux": n_panneaux, "n_batteries": n_batteries, "type_menage": type_menage,
            "duree_annees": duree_annees,
            "surface_max_m2": surface_max,
            "fichier_consommation": selecteur_conso.valeur(), "fichier_meteo": selecteur_meteo.valeur(),
        }
        lancer_en_arriere_plan(
            fenetre_racine, bouton, barre,
            fonction_calcul=lambda: executer_evaluation(config),
            fonction_affichage=lambda r: afficher_resultat_evaluation(zone, r),
        )

    bouton = ctk.CTkButton(onglet, text="▶  Lancer l'évaluation", command=lancer)
    bouton.grid(row=2, column=0, sticky="w", padx=15, pady=(8, 8))
    barre.grid(row=3, column=0, sticky="ew", padx=15, pady=(0, 8))
    zone.grid(row=4, column=0, sticky="nsew", padx=15, pady=(0, 15))


# --- Onglet 3 : Dimensionnement d'une nouvelle installation -----------------

def construire_onglet_dimensionnement(onglet, fenetre_racine):
    onglet.grid_columnconfigure(0, weight=1)
    onglet.grid_rowconfigure(5, weight=1)

    cadre_saisie = ctk.CTkFrame(onglet)
    cadre_saisie.grid(row=0, column=0, sticky="ew", padx=15, pady=(15, 8))
    ctk.CTkLabel(cadre_saisie, text="Contraintes de dimensionnement", font=ctk.CTkFont(weight="bold")).grid(
        row=0, column=0, columnspan=2, sticky="w", padx=10, pady=(10, 6))

    ctk.CTkLabel(cadre_saisie, text="Budget maximal").grid(row=1, column=0, sticky="w", padx=10, pady=4)
    champ_budget = ctk.CTkEntry(cadre_saisie, width=150)
    champ_budget.insert(0, "10000000")
    champ_budget.grid(row=1, column=1, sticky="w", padx=10, pady=4)

    ctk.CTkLabel(cadre_saisie, text="Surface maximale disponible (m²)").grid(row=2, column=0, sticky="w", padx=10, pady=4)
    champ_surface = ctk.CTkEntry(cadre_saisie, width=150)
    champ_surface.grid(row=2, column=1, sticky="w", padx=10, pady=4)

    ctk.CTkLabel(cadre_saisie, text="Type de ménage / site").grid(row=3, column=0, sticky="w", padx=10, pady=(4, 10))
    valeurs_menage = [f"{cle} — {v['libelle']}" for cle, v in TYPES_MENAGE.items()]
    combo_menage = ctk.CTkComboBox(cadre_saisie, values=valeurs_menage, width=280, state="readonly")
    combo_menage.set(f"{TYPE_MENAGE_PAR_DEFAUT} — {TYPES_MENAGE[TYPE_MENAGE_PAR_DEFAUT]['libelle']}")
    combo_menage.grid(row=3, column=1, sticky="w", padx=10, pady=(4, 10))

    cadre_fichiers = ctk.CTkFrame(onglet)
    cadre_fichiers.grid(row=1, column=0, sticky="ew", padx=15, pady=8)
    ctk.CTkLabel(cadre_fichiers, text="Données", font=ctk.CTkFont(weight="bold")).pack(anchor="w", padx=10, pady=(10, 6))
    selecteur_catalogue = SelecteurFichier(cadre_fichiers, "Catalogue matériel (CSV)", obligatoire=True)
    selecteur_catalogue.pack(fill="x", padx=10, pady=4)
    selecteur_conso = SelecteurFichier(cadre_fichiers, "Consommation mensuelle (CSV)")
    selecteur_conso.pack(fill="x", padx=10, pady=4)
    selecteur_meteo = SelecteurFichier(cadre_fichiers, "Ensoleillement journalier (CSV)")
    selecteur_meteo.pack(fill="x", padx=10, pady=(4, 10))

    ctk.CTkLabel(
        onglet, text=("Le catalogue matériel (*) est obligatoire : il doit contenir au moins un panneau, "
                      "une batterie ET un onduleur. Contrairement à la consommation et à la météo, il "
                      "n'existe pas de valeur par défaut universelle pour des prix de matériel."),
        text_color=COULEUR_AVERTISSEMENT, wraplength=650, justify="left",
    ).grid(row=2, column=0, sticky="w", padx=15, pady=(0, 8))

    barre = ctk.CTkProgressBar(onglet, mode="indeterminate")
    zone = ZoneResultats(onglet)

    def lancer():
        try:
            budget = float(champ_budget.get())
        except ValueError:
            messagebox.showerror("Saisie invalide", "Le budget doit être un nombre.")
            return
        if not selecteur_catalogue.valeur():
            messagebox.showerror("Fichier manquant", "Le catalogue matériel (CSV) est obligatoire.")
            return

        type_menage = combo_menage.get().split(" — ")[0] if combo_menage.get() else None
        surface_max = champ_surface.get().strip()
        config = {
            "budget_max": budget, "surface_max_m2": surface_max,
            "type_menage": type_menage,
            "fichier_catalogue": selecteur_catalogue.valeur(),
            "fichier_consommation": selecteur_conso.valeur(), "fichier_meteo": selecteur_meteo.valeur(),
        }
        lancer_en_arriere_plan(
            fenetre_racine, bouton, barre,
            fonction_calcul=lambda: executer_dimensionnement(config),
            fonction_affichage=lambda r: afficher_resultat_dimensionnement(zone, r),
        )

    bouton = ctk.CTkButton(onglet, text="▶  Lancer le dimensionnement (peut prendre 30 à 60 s)", command=lancer)
    bouton.grid(row=3, column=0, sticky="w", padx=15, pady=(0, 8))
    barre.grid(row=4, column=0, sticky="ew", padx=15, pady=(0, 8))
    zone.grid(row=5, column=0, sticky="nsew", padx=15, pady=(0, 15))


def preparer_donnees_monte_carlo():
    """Construit une analyse de robustesse en répétant plusieurs fois la procédure de dimensionnement."""
    chemin_catalogue = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "temp_catalogue_test.csv"))
    if not os.path.isfile(chemin_catalogue):
        raise DualSolarStatError(
            "Catalogue de test introuvable pour l'analyse de robustesse Monte Carlo. "
            "Placez le fichier temp_catalogue_test.csv à la racine du projet."
        )

    config = {
        "fichier_catalogue": chemin_catalogue,
        "budget_max": 15_000_000,
        "surface_max_m2": 200.0,
        "type_menage": TYPE_MENAGE_PAR_DEFAUT,
        "fichier_consommation": None,
        "fichier_meteo": None,
        "n_repetitions_par_config": 180,
    }

    def qualifier_dispersion(valeurs):
        q1, q3 = np.percentile(valeurs, [25, 75])
        iqr = q3 - q1
        if iqr == 0:
            return "très stable"
        if iqr <= 1:
            return "relativement stable"
        if iqr <= 2:
            return "modérément dispersé"
        return "fortement dispersé"

    repetitions = 20
    panneaux = []
    batteries = []
    erreurs = []
    for _ in range(repetitions):
        resultat = executer_dimensionnement(config)
        if not resultat.get("succes"):
            erreurs.append(resultat.get("erreur", "Erreur inconnue."))
            continue
        recommandation = resultat["configuration_recommandee"]
        panneaux.append(recommandation["n_panneaux"])
        batteries.append(recommandation["n_batteries"])

    if not panneaux:
        raise DualSolarStatError(
            "Impossible de générer l'analyse de robustesse : toutes les exécutions de dimensionnement ont échoué. "
            f"Erreurs rencontrées : {' ; '.join(erreurs[:3])}"
        )

    texte_resume = (
        f"Analyse sur {len(panneaux)} répétitions : recommandation panneaux {qualifier_dispersion(panneaux)} "
        f"(médiane {int(np.median(panneaux))}, étendue {min(panneaux)}–{max(panneaux)}),"
        f" recommandations batteries {qualifier_dispersion(batteries)} "
        f"(médiane {int(np.median(batteries))}, étendue {min(batteries)}–{max(batteries)})."
    )

    return {
        "serie_panneaux": np.asarray(panneaux, dtype=int),
        "serie_batteries": np.asarray(batteries, dtype=int),
        "summary": texte_resume,
    }


def _enregistrer_figure_temporaire(fig):
    chemin = os.path.join(tempfile.gettempdir(), f"dualsolarstat_graph_{uuid.uuid4().hex}.png")
    fig.savefig(chemin, dpi=150)
    return chemin


def afficher_graphe_monte_carlo(frame, donnees, image_state=None):
    """Affiche un indicateur de robustesse dans un frame Tkinter."""
    for widget in frame.winfo_children():
        widget.destroy()

    fig = Figure(figsize=(8.2, 5.2), dpi=100)
    ax = fig.add_subplot(111)

    if isinstance(donnees, dict) and "serie_panneaux" in donnees and "serie_batteries" in donnees:
        series = [donnees["serie_panneaux"], donnees["serie_batteries"]]
        labels = ["Panneaux", "Batteries"]
        bp = ax.boxplot(series, labels=labels, patch_artist=True, showmeans=True, meanline=True)
        for patch, color in zip(bp["boxes"], ["#4C78A8", "#F58518"]):
            patch.set_facecolor(color)
            patch.set_alpha(0.6)
        for i, serie in enumerate(series, start=1):
            x = np.full(len(serie), i) + (np.random.rand(len(serie)) - 0.5) * 0.15
            ax.scatter(x, serie, color="#222222", alpha=0.7, s=24, zorder=10)

        ax.set_title("Robustesse du dimensionnement : recommandations répétées")
        ax.set_xlabel("Type de composant")
        ax.set_ylabel("Nombre recommandé")
        ax.grid(True, alpha=0.25, axis="y")
        fig.tight_layout(rect=[0, 0.05, 1, 1])
        if donnees.get("summary"):
            fig.text(0.5, 0.01, donnees["summary"], ha="center", va="bottom", fontsize=9)
    else:
        iterations = np.arange(1, len(next(iter(donnees.values()))) + 1)
        for label, serie in donnees.items():
            ax.plot(iterations, serie, linewidth=1.8, label=label)
        ax.set_title("Convergence Monte Carlo")
        ax.set_xlabel("Nombre de simulations")
        ax.set_ylabel("Fiabilité moyenne cumulative")
        ax.set_ylim(0, 1.02)
        ax.grid(True, alpha=0.3)
        ax.legend(loc="best")
        fig.tight_layout()

    if image_state is not None:
        image_state["path"] = _enregistrer_figure_temporaire(fig)

    canvas = FigureCanvasTkAgg(fig, master=frame)
    canvas.draw()
    canvas.get_tk_widget().pack(fill="both", expand=True)


def construire_onglet_monte_carlo(onglet, fenetre_racine):
    onglet.grid_columnconfigure(0, weight=1)
    onglet.grid_rowconfigure(1, weight=1)

    ctk.CTkLabel(
        onglet,
        text=("Analyse de robustesse du dimensionnement : cette application répète plusieurs fois la procédure "
              "de dimensionnement sur le même catalogue et compare les nombres de panneaux et de batteries recommandés. "
              "Le boxplot met en évidence si les propositions se regroupent autour d'une solution stable ou si elles sont dispersées."),
        justify="left",
        anchor="w",
        wraplength=700,
    ).grid(row=0, column=0, sticky="w", padx=15, pady=(15, 8))

    barre = ctk.CTkProgressBar(onglet, mode="indeterminate")
    barre.grid(row=1, column=0, sticky="ew", padx=15, pady=(0, 8))

    cadre_graph = ctk.CTkFrame(onglet)
    cadre_graph.grid(row=2, column=0, sticky="nsew", padx=15, pady=(0, 15))
    cadre_graph.grid_columnconfigure(0, weight=1)
    cadre_graph.grid_rowconfigure(0, weight=1)

    image_state = {"path": None}

    def exporter_monte_carlo():
        if image_state["path"] is None or not os.path.isfile(image_state["path"]):
            messagebox.showerror("Aucun graphe", "Générez d'abord l'analyse de robustesse avant d'exporter le rapport Word.")
            return
        texte = (
            "Rapport de robustesse DualSolarStat\n\n"
            "Ce document contient une analyse des recommandations de dimensionnement obtenues sur plusieurs exécutions indépendantes.\n"
            "Le boxplot montre la dispersion des nombres de panneaux et de batteries recommandés, ce qui permet d'estimer "
            "la stabilité de la solution optimale face au bruit de simulation.\n"
        )
        exporter_resultat(None, texte, image_path=image_state["path"])

    def lancer():
        lancer_en_arriere_plan(
            fenetre_racine,
            bouton,
            barre,
            fonction_calcul=preparer_donnees_monte_carlo,
            fonction_affichage=lambda donnees: [
                afficher_graphe_monte_carlo(cadre_graph, donnees, image_state=image_state),
                bouton_exporter.configure(state="normal"),
            ],
        )

    bouton = ctk.CTkButton(onglet, text="▶  Générer l'analyse de robustesse", command=lancer)
    bouton.grid(row=3, column=0, sticky="w", padx=15, pady=(0, 8))

    bouton_exporter = ctk.CTkButton(
        onglet,
        text="💾 Exporter le rapport Word",
        width=220,
        command=exporter_monte_carlo,
        state="disabled",
    )
    bouton_exporter.grid(row=3, column=0, sticky="e", padx=15, pady=(0, 8))


# ============================================================================
# ==== SECTION C : POINT D'ENTREE ===========================================
# ============================================================================

class ApplicationDualSolarStat(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("DualSolarStat — Dimensionnement solaire probabiliste")
        self.geometry("880x760")
        self.minsize(760, 640)

        entete = ctk.CTkFrame(self, fg_color="transparent")
        entete.pack(fill="x", padx=20, pady=(15, 0))
        ctk.CTkLabel(entete, text="☀ DualSolarStat", font=ctk.CTkFont(size=22, weight="bold")).pack(anchor="w")
        ctk.CTkLabel(entete, text="Comparaison du dimensionnement solaire traditionnel et probabiliste",
                     font=ctk.CTkFont(size=12), text_color=("#555555", "#AAAAAA")).pack(anchor="w")

        self.theme_var = ctk.StringVar(value="Clair")
        self.menu_theme = ctk.CTkOptionMenu(
            entete,
            values=["Système", "Clair", "Sombre"],
            variable=self.theme_var,
            command=self._changer_theme,
            width=110,
        )
        self.menu_theme.pack(anchor="e", pady=(6, 0))

        ctk.CTkButton(
            entete,
            text="✕ Fermer",
            width=90,
            command=self.destroy,
            fg_color="#E5484D",
            hover_color="#C93B44",
        ).pack(anchor="e", pady=(6, 0), padx=(8, 0))

        onglets = ctk.CTkTabview(self)
        onglets.pack(fill="both", expand=True, padx=15, pady=15)

        nom_onglet_1 = "1. Évaluer une installation"
        nom_onglet_2 = "2. Dimensionner une installation"
        nom_onglet_3 = "3. Robustesse"
        onglets.add(nom_onglet_1)
        onglets.add(nom_onglet_2)
        onglets.add(nom_onglet_3)

        construire_onglet_evaluation(onglets.tab(nom_onglet_1), self)
        construire_onglet_dimensionnement(onglets.tab(nom_onglet_2), self)
        construire_onglet_monte_carlo(onglets.tab(nom_onglet_3), self)

    def _changer_theme(self, choix):
        modes = {
            "Système": "system",
            "Clair": "light",
            "Sombre": "dark",
        }
        ctk.set_appearance_mode(modes.get(choix, "light"))


def main():
    app = ApplicationDualSolarStat()
    app.mainloop()


if __name__ == "__main__":
    main()